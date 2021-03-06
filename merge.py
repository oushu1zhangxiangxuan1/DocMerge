from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.shared import Pt


def main():
    doc_ch = Document('/Users/johnsaxon/Documents/原文.docx')

    doc_en = Document('/Users/johnsaxon/Documents/翻译.docx')

    doc_out = Document()

    styles = doc_out.styles
    s_ch = styles.add_style("style_ch", WD_STYLE_TYPE.PARAGRAPH)
    s_ch.font.name = '宋体'
    s_ch.font.size = Pt(14)
    s_ch.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    s_en = styles.add_style("style_en", WD_STYLE_TYPE.PARAGRAPH)
    s_en.font.name = 'Times New Roman'
    s_en.font.size = Pt(14)
    s_en.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # print(doc_ch.styles['Body Text'].font.name)
    print(doc_en.styles)
    for s in doc_en.styles:
        print(s)
        print(s.name)

    # return

    print("ch段落数：", len(doc_ch.paragraphs))

    print("en段落数：", len(doc_en.paragraphs))

    # i=0
    # for p in doc_ch.paragraphs:
    #     i+=1
    #     print(p.text)
    #     if i>10:
    #         break
    len_ch = len(doc_ch.paragraphs)

    len_en = len(doc_en.paragraphs)

    len_out = 0
    if len_ch < len_en:
        print("中文行数较少")
        len_out = len_ch
    elif len_en < len_ch:
        print("英文行数较少")
        len_out = len_en
    else:
        len_out = len_en

    # print(doc_ch.paragraphs[1].paragraph_format.font)
    print(doc_ch.paragraphs[1].style)
    for i in range(len_out):
        # print(doc_ch.paragraphs[i].text)
        # print(doc_en.paragraphs[i].text)
        doc_out.add_paragraph(
            doc_ch.paragraphs[i].text, s_ch)
        # t.style = doc_ch.paragraph_format.font
        doc_out.add_paragraph(
            doc_en.paragraphs[i].text, s_en)

    doc_out.save("合并.docx")


if '__main__' == __name__:
    main()
