from docx import Document
import os
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.shared import Pt


def merge(ch, en, out):
    doc_ch = Document(ch)

    doc_en = Document(en)

    doc_out = Document()

    styles = doc_out.styles
    s_ch = styles.add_style("style_ch", WD_STYLE_TYPE.PARAGRAPH)
    s_ch.font.name = '宋体'
    s_ch.font.size = Pt(14)
    s_ch.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    s_en = styles.add_style("style_en", WD_STYLE_TYPE.PARAGRAPH)
    s_en.font.name = 'Times New Roman'
    s_en.font.size = Pt(14)
    s_en.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    len_ch = len(doc_ch.paragraphs)

    len_en = len(doc_en.paragraphs)

    len_out = 0
    if len_ch < len_en:
        print(ch)
        print(en)
        print("中文行数较少: ", len_ch-len_en)
        print("\n\n")
        len_out = len_ch
    elif len_en < len_ch:
        print(ch)
        print(en)
        print("英文行数较少:", len_ch-len_en)
        print("\n\n")
        len_out = len_en
    else:
        len_out = len_en

    for i in range(len_out):
        doc_out.add_paragraph(
            doc_ch.paragraphs[i].text, s_ch)
        doc_out.add_paragraph(
            doc_en.paragraphs[i].text, s_en)

    doc_out.save(out)


def walk(doc_path, out_path):
    for root, dirs, files in os.walk(doc_path):
        for d in dirs:
            for sub_root, _, sub_files in os.walk(os.path.join(root, d)):
                if len(sub_files) != 2:
                    print(sub_files, " failed.")
                    continue
                if len(sub_files[0]) > len(sub_files[1]):
                    doc_ch = os.path.join(sub_root, sub_files[0])
                    doc_out = os.path.join(out_path, sub_files[0])
                    doc_en = os.path.join(sub_root, sub_files[1])
                else:
                    doc_ch = os.path.join(sub_root, sub_files[1])
                    doc_out = os.path.join(out_path, sub_files[1])
                    doc_en = os.path.join(sub_root, sub_files[0])

                merge(doc_ch, doc_en, doc_out)


def main():
    root_doc_path = "/Users/johnsaxon/Documents/transMerge"
    out_path = "/Users/johnsaxon/Documents/transMerge/merge"
    for root, dirs, files in os.walk(root_doc_path):
        # print(root)
        # print(dirs)
        # print(files)
        for d in dirs:
            if d == "merge":
                print("jump merge")
                continue
            walk(os.path.join(root, d), out_path)


if '__main__' == __name__:
    main()
